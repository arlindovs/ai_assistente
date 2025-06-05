import os
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import textract
from llama_index import Document, GPTVectorStoreIndex, LLMPredictor, ServiceContext
from llama_index.vector_stores.chroma import ChromaVectorStore
import chromadb

# --------- CONFIGURAÇÕES ---------
PASTA_ARQUIVOS = "./arquivos_para_indexar"  # Altere para o seu caminho real!
CHROMA_PATH = "./chroma_db"
COLLECTION = "documentos_texto"

llm_predictor = LLMPredictor(llm=None)
service_context = ServiceContext.from_defaults(
    llm_predictor=llm_predictor,
    embed_model="local"
)

def extrair_textos_de_arquivos(pasta):
    docs = []
    for filename in os.listdir(pasta):
        caminho = os.path.join(pasta, filename)
        if filename.endswith('.txt'):
            with open(caminho, 'r', encoding='utf-8') as f:
                texto = f.read()
                docs.append((filename, texto))
        elif filename.endswith('.docx'):
            docx_doc = DocxDocument(caminho)
            texto = "\n".join([p.text for p in docx_doc.paragraphs])
            docs.append((filename, texto))
        elif filename.endswith('.doc'):
            try:
                texto = textract.process(caminho).decode('utf-8', errors='ignore')
                docs.append((filename, texto))
            except Exception as e:
                print(f"Erro ao ler DOC {filename}: {e}")
        elif filename.endswith('.pdf'):
            try:
                pdf = PdfReader(caminho)
                texto = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        texto += page_text + "\n"
                docs.append((filename, texto))
            except Exception as e:
                print(f"Erro ao ler PDF {filename}: {e}")
    return docs

def criar_documentos_llama(textos):
    return [Document(text=f"[Arquivo: {nome}]\n{txt}") for nome, txt in textos]

def indexar_arquivos_em_chroma(documentos_llama):
    db = chromadb.PersistentClient(path=CHROMA_PATH)
    collection = db.get_or_create_collection(COLLECTION)
    # NÃO limpa a collection — apenas acumula!
    chroma_store = ChromaVectorStore(chroma_collection=collection)
    index = GPTVectorStoreIndex.from_documents(
        documentos_llama,
        vector_store=chroma_store,
        service_context=service_context
    )
    print("Indexação concluída.")

if __name__ == "__main__":
    print(f"Lendo arquivos da pasta: {PASTA_ARQUIVOS}")
    textos = extrair_textos_de_arquivos(PASTA_ARQUIVOS)
    print(f"{len(textos)} arquivos lidos.")
    documentos_llama = criar_documentos_llama(textos)
    print("Indexando arquivos no ChromaDB (sem limpeza)...")
    indexar_arquivos_em_chroma(documentos_llama)
    print("Script finalizado.")
